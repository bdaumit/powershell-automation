# export_xanthus.py
from docx import Document
from pathlib import Path
import os
import sys

def docx_to_text(docx_path: Path) -> str:
    doc = Document(docx_path)
    lines = []

    # Paragraphs
    for p in doc.paragraphs:
        lines.append(p.text)

    # Tables (optional but handy if you ever use them)
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [c.text for c in row.cells]
            # join cells with tabs; change to ' | ' if you prefer
            lines.append("\t".join(cells))

    # Normalize: strip trailing spaces on each line, collapse extra blank lines a bit
    text = "\n".join(line.rstrip() for line in lines)
    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")
    return text

def main(docx_path_str: str, txt_path_str: str):
    docx_path = Path(docx_path_str).expanduser().resolve()
    txt_path  = Path(txt_path_str).expanduser().resolve()

    if not docx_path.exists():
        print(f"ERROR: .docx not found: {docx_path}")
        sys.exit(1)

    text = docx_to_text(docx_path)

    # Write atomically: write to temp then replace
    tmp_path = txt_path.with_suffix(".txt.tmp")
    tmp_path.parent.mkdir(parents=True, exist_ok=True)
    with open(tmp_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(text)
    os.replace(tmp_path, txt_path)  # atomic on Windows & Unix

    # Simple word count feedback
    words = len(text.split())
    print(f"✓ Updated {txt_path.name} ({words} words) from {docx_path.name}")

if __name__ == "__main__":
    # Usage:
    #   python export_xanthus.py "C:\path\Xanthus.docx" "C:\path\Xanthus.txt"
    # If you omit args, it defaults to Xanthus.* next to the script.
    if len(sys.argv) >= 3:
        main(sys.argv[1], sys.argv[2])
    else:
        base = Path(__file__).parent / "Xanthus"
        main(str(base.with_suffix(".docx")), str(base.with_suffix(".txt")))
